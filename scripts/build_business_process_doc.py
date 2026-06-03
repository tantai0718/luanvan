from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = Path(r"D:\thuctap\chonongsan\cho_nong_san_FINAL\project\Mo_ta_nghiep_vu_website_cho_nong_san.docx")


PRIMARY = RGBColor(0x2E, 0x74, 0xB5)
SECONDARY = RGBColor(0x1F, 0x4D, 0x78)
TEXT = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
LIGHT = RGBColor(0xF2, 0xF4, 0xF7)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "D8DDE6")
        borders.append(elem)
    tbl_pr.append(borders)


def style_run(run, *, bold=False, italic=False, size=11, color=TEXT, font="Calibri"):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)


def add_paragraph(doc, text="", style=None, align=None, before=0, after=6, line=1.1):
    p = doc.add_paragraph(style=style)
    if text:
      run = p.add_run(text)
      style_run(run)
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align:
        p.alignment = align
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.18)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    style_run(run)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    style_run(run)
    return p


def create_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(24)
    title.font.color.rgb = PRIMARY
    title.font.bold = True
    title.paragraph_format.space_after = Pt(8)

    for name, size, color in (
        ("Heading 1", 16, PRIMARY),
        ("Heading 2", 13, PRIMARY),
        ("Heading 3", 12, SECONDARY),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    if "Callout" not in [s.name for s in doc.styles]:
        callout = doc.styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
        callout.base_style = doc.styles["Normal"]
        callout.font.name = "Calibri"
        callout._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        callout._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        callout.font.size = Pt(11)
        callout.font.color.rgb = TEXT


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("Website chợ nông sản - Mô tả nghiệp vụ")
    style_run(header_run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Tài liệu mô tả các nghiệp vụ đã được triển khai trong hệ thống.")
    style_run(footer_run, size=9, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    r = p.add_run("MÔ TẢ NGHIỆP VỤ\nHỆ THỐNG WEBSITE CHỢ NÔNG SẢN")
    style_run(r, bold=True, size=22, color=PRIMARY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run("Tài liệu này tổng hợp các luồng nghiệp vụ hiện đã có trên website, chỉ bám theo chức năng đã được triển khai.")
    style_run(r2, size=11, color=MUTED)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.1)
    table.columns[1].width = Inches(4.2)
    set_table_borders(table)
    rows = [
        ("Tên tài liệu", "Mô tả nghiệp vụ hệ thống website chợ nông sản"),
        ("Phạm vi", "Luồng người mua, nông dân, quản trị viên và kho hàng"),
        ("Nguồn đối chiếu", "Frontend ReactJS, backend Node.js/ExpressJS và cơ sở dữ liệu MySQL của dự án hiện tại"),
        ("Ngày lập", datetime.now().strftime("%d/%m/%Y")),
    ]
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).width = Inches(2.1)
        table.cell(i, 1).width = Inches(4.2)
        for j, text in enumerate((label, value)):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            run = cell.paragraphs[0].add_run(text)
            style_run(run, bold=j == 0)
            if j == 0:
                set_cell_shading(cell, "F2F4F7")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)


def add_section_title(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    style_run(run, bold=True, size=16, color=PRIMARY)


def add_subtitle(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, bold=True, size=13, color=PRIMARY)


def add_module_table(doc):
    add_subtitle(doc, "1.1. Nhóm chức năng hiện có trên hệ thống")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.9), Inches(1.8), Inches(2.8)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    set_table_borders(table)

    headers = ["Nhóm nghiệp vụ", "Vai trò sử dụng", "Mô tả ngắn"]
    for idx, header_text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "E8EEF5")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = cell.paragraphs[0].add_run(header_text)
        style_run(run, bold=True)

    rows = [
        ("Xác thực", "Người mua, nông dân, admin", "Đăng ký, đăng nhập, xem hồ sơ và đổi mật khẩu."),
        ("Sản phẩm", "Người mua, nông dân, admin", "Xem danh sách, xem chi tiết, quản lý sản phẩm, ẩn hiện và cập nhật thông tin."),
        ("Mua hàng", "Người mua", "Giỏ hàng, đơn hàng thường, hủy đơn, theo dõi đơn hàng."),
        ("Đặt trước", "Người mua, admin", "Tạo đơn đặt trước với ngày giao dự kiến cho đợt hàng tiếp theo."),
        ("Giao định kỳ", "Người mua, nông dân, admin", "Đăng ký giao lặp lại theo tuần, hai tuần hoặc tháng."),
        ("Kho hàng", "Admin", "Nhập xuất kho, tồn kho, vị trí lưu, hạn sử dụng, cảnh báo kho."),
        ("Quản trị hệ thống", "Admin", "Tài khoản, danh mục, đơn hàng, kiểm soát nông dân và dữ liệu vận hành."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].width = widths[idx]
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cells[idx].paragraphs[0].paragraph_format.space_after = Pt(0)
            run = cells[idx].paragraphs[0].add_run(text)
            style_run(run)


def add_flow_table(doc, title, steps):
    add_subtitle(doc, title)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(0.8), Inches(2.0), Inches(3.7)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    set_table_borders(table)

    headers = ["Bước", "Thao tác", "Kết quả nghiệp vụ"]
    for idx, header_text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "F2F4F7")
        run = cell.paragraphs[0].add_run(header_text)
        style_run(run, bold=True)

    for number, action, result in steps:
        cells = table.add_row().cells
        values = (str(number), action, result)
        for idx, text in enumerate(values):
            cells[idx].width = widths[idx]
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cells[idx].paragraphs[0].paragraph_format.space_after = Pt(0)
            run = cells[idx].paragraphs[0].add_run(text)
            style_run(run)


def build():
    doc = Document()
    create_styles(doc)
    configure_page(doc)
    add_cover(doc)

    add_section_title(doc, "1. Tổng quan tài liệu")
    add_paragraph(
        doc,
        "Tài liệu này mô tả các nghiệp vụ đang vận hành trong hệ thống website chợ nông sản. Nội dung được tổng hợp trực tiếp từ mã nguồn hiện tại của dự án, bao gồm giao diện frontend, các API backend và cấu trúc dữ liệu liên quan. Mục tiêu của tài liệu là giúp người đọc hiểu rõ vai trò người dùng, các bước xử lý nghiệp vụ và mối liên hệ giữa những chức năng đã được triển khai."
    )
    add_paragraph(
        doc,
        "Phạm vi trình bày chỉ bao gồm những chức năng đã có trong hệ thống tại thời điểm lập tài liệu. Tài liệu không mô tả các tính năng dự kiến trong tương lai hoặc những ý tưởng chưa được hiện thực hóa trong web."
    )
    add_module_table(doc)

    add_section_title(doc, "2. Tác nhân và phạm vi sử dụng")
    add_paragraph(doc, "Hệ thống hiện có ba vai trò người dùng chính, mỗi vai trò đảm nhận một nhóm nghiệp vụ riêng.")
    add_bullet(doc, "Người mua: xem sản phẩm, thêm vào giỏ, tạo đơn hàng thường, đặt trước, đăng ký giao định kỳ, theo dõi đơn và gửi đánh giá sau khi giao hàng.")
    add_bullet(doc, "Nông dân: quản lý hồ sơ nông trại, thêm sửa sản phẩm, theo dõi đơn hàng liên quan đến nông trại và xem các đăng ký giao định kỳ phát sinh từ sản phẩm của mình.")
    add_bullet(doc, "Quản trị viên: quản lý toàn bộ tài khoản, danh mục, sản phẩm, đơn hàng, đăng ký định kỳ và nghiệp vụ kho hàng.")
    add_paragraph(doc, "Bên cạnh ba vai trò trên, hệ thống còn có lớp dữ liệu kho hàng để hỗ trợ nhập xuất, cảnh báo tồn kho và theo dõi vị trí lưu trữ sản phẩm.")

    add_section_title(doc, "3. Nghiệp vụ phía người mua")
    add_paragraph(doc, "Nhóm nghiệp vụ người mua là phần trung tâm của website, bao gồm toàn bộ luồng từ tiếp cận sản phẩm, đặt hàng cho đến theo dõi và đánh giá.")
    add_flow_table(doc, "3.1. Luồng xem và lựa chọn sản phẩm", [
        (1, "Truy cập trang chủ hoặc trang sản phẩm", "Hệ thống tải danh sách sản phẩm công khai, hỗ trợ lọc theo danh mục, giá, khu vực, tình trạng còn hàng và cách sắp xếp."),
        (2, "Mở trang chi tiết sản phẩm", "Người mua xem được ảnh, mô tả, giá bán, đơn vị, tồn kho, thông tin nông trại, đánh giá và sản phẩm liên quan."),
        (3, "Chọn số lượng", "Hệ thống kiểm soát số lượng không vượt quá tồn kho hiện tại của sản phẩm."),
        (4, "Thêm vào giỏ hoặc mua ngay", "Nếu chưa đăng nhập thì chuyển đến màn đăng nhập; nếu hợp lệ thì thêm sản phẩm vào giỏ hàng."),
    ])
    add_flow_table(doc, "3.2. Luồng mua hàng thông thường", [
        (1, "Mở giỏ hàng", "Người mua xem được các sản phẩm đã chọn, số lượng, đơn giá và tổng tiền tạm tính."),
        (2, "Cập nhật giỏ hàng", "Có thể tăng giảm số lượng, xóa từng sản phẩm hoặc xóa toàn bộ giỏ."),
        (3, "Nhập địa chỉ giao hàng và ghi chú", "Thông tin giao nhận được gửi kèm khi tạo đơn hàng."),
        (4, "Chọn phương thức thanh toán", "Hệ thống ghi nhận hình thức thanh toán hiện có như tiền mặt hoặc VNPay ở mức lựa chọn giao diện."),
        (5, "Tạo đơn hàng", "Backend sinh đơn hàng, chi tiết đơn hàng, cập nhật tồn kho và trả lại thông tin đơn cho người mua."),
        (6, "Theo dõi trạng thái đơn", "Người mua có thể xem danh sách đơn, chi tiết đơn và hủy đơn trong các trạng thái cho phép."),
    ])
    add_flow_table(doc, "3.3. Luồng đặt trước nông sản", [
        (1, "Mở khối Đặt trước ở trang chi tiết sản phẩm", "Người mua nhập số lượng, ngày giao dự kiến, địa chỉ và ghi chú riêng cho đơn đặt trước."),
        (2, "Xác nhận tạo đơn đặt trước", "Hệ thống tạo đơn với loại đơn là đặt trước và lưu ngày giao dự kiến."),
        (3, "Theo dõi đơn đặt trước", "Đơn đặt trước xuất hiện cùng hệ thống đơn hàng, giúp người mua quản lý tập trung."),
    ])
    add_flow_table(doc, "3.4. Luồng đăng ký giao định kỳ", [
        (1, "Mở khối Giao định kỳ ở trang chi tiết sản phẩm", "Người mua nhập số lượng mỗi kỳ, ngày bắt đầu, tần suất giao và số kỳ giao."),
        (2, "Gửi đăng ký", "Hệ thống tạo bản ghi đăng ký giao định kỳ, xác định ngày giao tiếp theo và trạng thái hoạt động."),
        (3, "Quản lý đăng ký", "Người mua xem được danh sách đăng ký định kỳ trong trang đơn hàng và có thể hủy khi cần."),
    ])
    add_flow_table(doc, "3.5. Luồng đánh giá sản phẩm", [
        (1, "Điều kiện đánh giá", "Người mua chỉ được đánh giá nếu đã có đơn hàng ở trạng thái đã giao và đơn đó chứa sản phẩm cần đánh giá."),
        (2, "Gửi số sao và nội dung nhận xét", "Hệ thống lưu bản ghi đánh giá, gắn với sản phẩm, người mua và đơn hàng tương ứng."),
        (3, "Cập nhật điểm trung bình", "Sau khi lưu đánh giá, hệ thống tính lại điểm đánh giá và tổng số lượt đánh giá của sản phẩm."),
    ])

    add_section_title(doc, "4. Nghiệp vụ phía nông dân")
    add_paragraph(doc, "Vai trò nông dân được thiết kế để quản lý gian hàng và theo dõi hoạt động bán hàng gắn với nông trại.")
    add_bullet(doc, "Xem dashboard tổng hợp số sản phẩm, đơn hàng, doanh thu và trạng thái tồn kho.")
    add_bullet(doc, "Cập nhật hồ sơ nông trại gồm tên nông trại, khu vực, thông tin giới thiệu và dữ liệu liên hệ.")
    add_bullet(doc, "Thêm sản phẩm mới, chỉnh sửa thông tin, thay đổi trạng thái hiển thị, cập nhật giá bán, tồn kho tối thiểu và hình ảnh sản phẩm.")
    add_bullet(doc, "Theo dõi các đơn hàng phát sinh từ sản phẩm của mình và cập nhật trạng thái giao hàng theo tiến trình xử lý.")
    add_bullet(doc, "Xem danh sách đăng ký giao định kỳ liên quan đến nông trại để chuẩn bị nguồn cung theo chu kỳ.")

    add_section_title(doc, "5. Nghiệp vụ phía quản trị viên")
    add_paragraph(doc, "Quản trị viên là vai trò kiểm soát dữ liệu và vận hành toàn bộ hệ thống.")
    add_flow_table(doc, "5.1. Quản lý tài khoản và nông dân", [
        (1, "Xem danh sách tài khoản", "Admin truy cập danh sách tài khoản, lọc theo vai trò và trạng thái."),
        (2, "Khóa hoặc mở khóa tài khoản", "Trạng thái hoạt động của tài khoản được cập nhật trực tiếp trên hệ thống."),
        (3, "Duyệt nông dân", "Admin xác minh hồ sơ nông dân trước khi cho phép vận hành đầy đủ."),
        (4, "Xóa nông dân", "Khi xóa, hệ thống dọn tài khoản, sản phẩm, tồn kho, đánh giá và dữ liệu vận hành liên quan đến nông dân đó."),
    ])
    add_flow_table(doc, "5.2. Quản lý danh mục, sản phẩm và đơn hàng", [
        (1, "Quản lý danh mục", "Admin thêm, sửa, ẩn hiện danh mục và xem các sản phẩm thuộc từng danh mục."),
        (2, "Quản lý sản phẩm toàn hệ thống", "Admin có thể ẩn hiện, theo dõi điểm đánh giá và kiểm soát thông tin sản phẩm."),
        (3, "Quản lý đơn hàng", "Admin xem danh sách đơn thường và đơn đặt trước, cập nhật trạng thái và hỗ trợ xử lý hủy đơn."),
        (4, "Theo dõi đăng ký giao định kỳ", "Admin xem được toàn bộ các đăng ký giao lặp lại để phục vụ công tác vận hành."),
    ])
    add_bullet(doc, "Dashboard quản trị cung cấp cái nhìn tổng hợp về tài khoản, nông dân, sản phẩm, đơn hàng và kho.")
    add_bullet(doc, "Thông báo hệ thống hỗ trợ nhắc việc và phản ánh các biến động quan trọng theo tài khoản đăng nhập.")

    add_section_title(doc, "6. Nghiệp vụ kho hàng")
    add_paragraph(doc, "Kho hàng là phần mở rộng giúp hệ thống không chỉ xử lý giao dịch mua bán mà còn quản lý hàng hóa ở mức vận hành thực tế.")
    add_flow_table(doc, "6.1. Luồng nhập - xuất kho", [
        (1, "Tạo hóa đơn kho", "Admin chọn loại hóa đơn nhập hoặc xuất, kho hàng, sản phẩm, số lượng, đơn giá và thông tin ghi chú."),
        (2, "Khai báo vị trí lưu và hạn sử dụng", "Mỗi dòng hàng có thể gắn vị trí trong kho, ngày nhập và hạn sử dụng để theo dõi chi tiết hơn."),
        (3, "Xác nhận hóa đơn", "Khi xác nhận, hệ thống cập nhật tồn kho, lịch sử kho và kiểm tra cảnh báo tồn."),
        (4, "Hủy hóa đơn", "Nếu hóa đơn bị hủy, hệ thống giữ được dấu vết trạng thái để đối chiếu nghiệp vụ."),
    ])
    add_flow_table(doc, "6.2. Theo dõi tồn kho và cảnh báo", [
        (1, "Xem tồn kho theo kho", "Admin biết được mỗi sản phẩm đang nằm ở kho nào, số lượng còn bao nhiêu và vị trí cất giữ."),
        (2, "Cập nhật metadata tồn kho", "Có thể bổ sung vị trí, hạn sử dụng và ngày nhập cho từng dòng tồn kho."),
        (3, "Kiểm tra cảnh báo", "Hệ thống phát hiện các trường hợp sắp hết hàng hoặc thiếu thông tin kho cần bổ sung."),
        (4, "Đồng bộ với đơn hàng", "Khi đơn hàng tạo hoặc bị hủy, hệ thống hỗ trợ trừ hoặc hoàn lại tồn kho tương ứng theo dữ liệu kho."),
    ])

    add_section_title(doc, "7. Mối liên kết giữa các nghiệp vụ")
    add_paragraph(doc, "Các nghiệp vụ trong hệ thống không tách rời nhau mà được liên kết theo một chuỗi xử lý thống nhất. Sản phẩm là trung tâm của luồng mua hàng, đồng thời cũng là đối tượng chính trong kho, đánh giá, đăng ký giao định kỳ và các báo cáo quản trị.")
    add_numbered(doc, "Người mua phát sinh nhu cầu thông qua trang sản phẩm, giỏ hàng, đơn hàng thường, đặt trước hoặc giao định kỳ.")
    add_numbered(doc, "Nông dân và quản trị viên đảm bảo thông tin sản phẩm, giá bán, hình ảnh và trạng thái hiển thị luôn được cập nhật.")
    add_numbered(doc, "Kho hàng đảm nhận việc theo dõi số lượng thực tế, vị trí lưu trữ, hạn sử dụng và biến động nhập xuất.")
    add_numbered(doc, "Đơn hàng và đánh giá tạo ra lớp dữ liệu phản hồi giúp hoàn thiện chất lượng thông tin sản phẩm.")
    add_paragraph(doc, "Nhờ cấu trúc này, hệ thống có thể vận hành như một website chợ nông sản có kiểm soát, đồng thời vẫn đủ linh hoạt để mở rộng thêm các nghiệp vụ sâu hơn trong tương lai.")

    add_section_title(doc, "8. Kết luận")
    add_paragraph(doc, "Tài liệu mô tả nghiệp vụ cho thấy hệ thống hiện đã triển khai tương đối đầy đủ các luồng cốt lõi của một website chợ nông sản: xác thực người dùng, quản lý sản phẩm, mua hàng, đặt trước, giao định kỳ, đánh giá, quản trị và kho hàng. Tất cả nội dung trong tài liệu đều được viết theo chức năng hiện có trong web, không bổ sung các nghiệp vụ chưa được hiện thực hóa.")
    add_paragraph(doc, "Tài liệu này có thể được dùng làm nền tảng cho báo cáo đồ án, mô tả hệ thống hoặc tiếp tục mở rộng thành tài liệu nghiệp vụ chi tiết hơn ở giai đoạn luận văn.")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
